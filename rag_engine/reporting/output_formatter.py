from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    for line in markdown_text.splitlines():
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def write_text_file(content: str, output_path: str) -> str:
    Path(output_path).write_text(content, encoding='utf-8')
    return output_path
